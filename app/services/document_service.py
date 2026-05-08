"""
Document Service — handles file uploads, text extraction, chunking, and embedding.
Reuses existing ChromaDB + AIMemoryIndex infrastructure (source_type=document).
"""

import logging
import os
import re
import uuid
from datetime import datetime
from typing import List, Optional

from sqlalchemy.orm import Session

from app.ai.chroma_manager import delete_embedding, store_embeddings_batch
from app.models.ai_memory_index import AIMemoryIndex, SourceType
from app.models.uploaded_document import DocumentStatus, UploadedDocument

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "md"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
CHUNK_SIZE = 800
CHUNK_OVERLAP = 150
UPLOADS_ROOT = "uploads"


class DocumentService:

    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        return re.sub(r"[^\w.\-]", "_", filename)

    @staticmethod
    def _extract_text(file_path: str, file_type: str) -> str:
        if file_type == "pdf":
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
            return "\n\n".join(pages)

        if file_type == "docx":
            from docx import Document

            doc = Document(file_path)
            return "\n\n".join(
                p.text for p in doc.paragraphs if p.text.strip()
            )

        # txt / md
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
        """Sliding-window chunker that prefers splitting on paragraph/line breaks."""
        if not text.strip():
            return []

        # Normalize whitespace
        text = re.sub(r"\n{3,}", "\n\n", text).strip()

        chunks: List[str] = []
        start = 0
        length = len(text)

        while start < length:
            end = min(start + chunk_size, length)

            if end < length:
                # Try to find a natural break point (\n\n first, then \n, then space)
                for sep in ("\n\n", "\n", " "):
                    pos = text.rfind(sep, start, end)
                    if pos > start:
                        end = pos + len(sep)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap if end - overlap > start else end

        return chunks

    @staticmethod
    def upload_document(
        db: Session,
        project_id: int,
        user_id: int,
        filename: str,
        file_bytes: bytes,
        file_type: str,
    ) -> UploadedDocument:
        """
        Save file to disk and create a DB row (status=processing).
        Embedding is handled separately via process_document_embeddings().
        """
        if file_type not in ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {file_type}. Allowed: {ALLOWED_EXTENSIONS}")

        if len(file_bytes) > MAX_FILE_SIZE:
            raise ValueError(f"File exceeds maximum size of {MAX_FILE_SIZE // (1024 * 1024)} MB")

        project_dir = os.path.join(UPLOADS_ROOT, str(project_id))
        os.makedirs(project_dir, exist_ok=True)

        safe_name = DocumentService._sanitize_filename(filename)
        unique_name = f"{uuid.uuid4().hex}_{safe_name}"
        file_path = os.path.join(project_dir, unique_name)

        with open(file_path, "wb") as f:
            f.write(file_bytes)

        doc = UploadedDocument(
            project_id=project_id,
            uploaded_by=user_id,
            filename=filename,
            file_path=file_path,
            file_type=file_type,
            status=DocumentStatus.processing,
            chunk_count=0,
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        logger.info(f"Document saved: {file_path} (id={doc.id})")
        return doc

    @staticmethod
    def process_document_embeddings(doc_id: int) -> None:
        """
        Background task: extract text → chunk → batch-embed in ChromaDB.
        Opens its own DB session so it can run outside the request lifecycle.
        """
        from app.db.session import SessionLocal

        db: Session = SessionLocal()
        try:
            doc = db.query(UploadedDocument).filter(UploadedDocument.id == doc_id).first()
            if not doc:
                logger.error(f"Document {doc_id} not found for processing")
                return

            text = DocumentService._extract_text(doc.file_path, doc.file_type)
            chunks = DocumentService._chunk_text(text)

            if not chunks:
                logger.warning(f"Document {doc_id} produced no chunks — marking failed")
                doc.status = DocumentStatus.failed
                db.commit()
                return

            embedding_ids = [str(uuid.uuid4()) for _ in chunks]
            now_iso = datetime.utcnow().isoformat()

            metadatas = [
                {
                    "project_id": doc.project_id,
                    "source_type": "document",
                    "source_id": doc.id,
                    "filename": doc.filename,
                    "chunk_index": i,
                    "created_at": now_iso,
                }
                for i in range(len(chunks))
            ]

            # Single ChromaDB batch call — 10-50x faster than per-chunk calls
            store_embeddings_batch(
                embedding_ids=embedding_ids,
                texts=chunks,
                metadatas=metadatas,
            )

            # Persist AIMemoryIndex rows in one batch
            memory_rows = [
                AIMemoryIndex(
                    project_id=doc.project_id,
                    source_type=SourceType.document,
                    source_id=doc.id,
                    embedding_id=eid,
                )
                for eid in embedding_ids
            ]
            db.add_all(memory_rows)

            doc.status = DocumentStatus.ready
            doc.chunk_count = len(chunks)
            db.commit()

            logger.info(
                f"Document {doc_id} processed: {len(chunks)} chunks embedded "
                f"(project={doc.project_id})"
            )

            # Auto-trigger summary regeneration now that a new document is available.
            # process_document_embeddings runs in a BackgroundTasks thread pool (sync),
            # so we use the thread-safe helper that schedules via call_soon_threadsafe.
            try:
                from app.services.background_summary_generator import queue_summary_generation_from_thread
                queue_summary_generation_from_thread(doc.project_id, "document")
            except Exception as _e:
                logger.warning(f"Could not trigger summary generation: {_e}")

        except Exception as e:
            logger.error(f"Failed to process document {doc_id}: {e}")
            try:
                doc = db.query(UploadedDocument).filter(UploadedDocument.id == doc_id).first()
                if doc:
                    doc.status = DocumentStatus.failed
                    db.commit()
            except Exception:
                pass
        finally:
            db.close()

    @staticmethod
    def delete_document(db: Session, doc_id: int, project_id: int) -> None:
        """Delete document from disk, ChromaDB, and DB (all chunks + metadata)."""
        doc = (
            db.query(UploadedDocument)
            .filter(
                UploadedDocument.id == doc_id,
                UploadedDocument.project_id == project_id,
            )
            .first()
        )
        if not doc:
            raise ValueError(f"Document {doc_id} not found in project {project_id}")

        # Remove all ChromaDB embeddings for this document
        memory_rows = (
            db.query(AIMemoryIndex)
            .filter(
                AIMemoryIndex.source_type == SourceType.document,
                AIMemoryIndex.source_id == doc_id,
            )
            .all()
        )
        for row in memory_rows:
            delete_embedding(row.embedding_id)
            db.delete(row)

        # Remove file from disk
        if os.path.exists(doc.file_path):
            try:
                os.remove(doc.file_path)
            except OSError as e:
                logger.warning(f"Could not delete file {doc.file_path}: {e}")

        db.delete(doc)
        db.commit()
        logger.info(f"Document {doc_id} deleted (project={project_id})")

    @staticmethod
    def list_documents(
        db: Session, project_id: int
    ) -> List[UploadedDocument]:
        return (
            db.query(UploadedDocument)
            .filter(UploadedDocument.project_id == project_id)
            .order_by(UploadedDocument.created_at.desc())
            .all()
        )

    @staticmethod
    def get_document(
        db: Session, doc_id: int, project_id: int
    ) -> Optional[UploadedDocument]:
        return (
            db.query(UploadedDocument)
            .filter(
                UploadedDocument.id == doc_id,
                UploadedDocument.project_id == project_id,
            )
            .first()
        )
